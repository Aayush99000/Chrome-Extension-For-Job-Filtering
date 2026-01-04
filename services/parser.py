import PyPDF2
from docx import Document
from io import BytesIO

def parse_pdf(file):

    try:
        doc = Document(BytesIO(file.read()))
        text = " "
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Extract tables (important for resumes!)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"

        if not text.strip():
            raise ValueError("🤖 No text found in the PDF file.")

    except Exception as e:
        raise ValueError(f"🤖 Error parsing PDF file: {str(e)}")

def parse_docx(file):
    """Extract text from DOCX file"""
    try:
        doc = Document(BytesIO(file.read()))
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Extract tables (important for resumes!)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        if not text.strip():
            raise ValueError("DOCX appears to be empty")
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Error parsing DOCX: {str(e)}")

def extract_sections(resume_text):
    """
    Split resume into meaningful sections
    Useful for better embeddings
    """
    # Simple paragraph-based splitting
    sections = [s.strip() for s in resume_text.split('\n\n') if s.strip()]
    
    # Filter out very short sections (less than 20 characters)
    sections = [s for s in sections if len(s) > 20]
    
    return sections

def parse_resume(file):
    filename =file.filename.lower()

    try:
        if filename.endswith('.pdf'):
            return parse_pdf(file)
        elif filename.endswith('.docx'):
            return parse_docx(file)
        elif filename.endswith('.doc'):
            return parse_docx(file)
        else:
            raise ValueError("🤖 Unsupported file format. Please upload a PDF or DOCX file.")
        
    except Exception as e:
        raise ValueError(f"🤖 Error parsing resume: {str(e)}")